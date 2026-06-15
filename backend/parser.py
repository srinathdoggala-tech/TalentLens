import io
import fitz  # PyMuPDF
import pdfplumber
import docx

def parse_pdf(file_bytes: bytes) -> str:
    """
    Parses PDF bytes using PyMuPDF (fitz) with a fallback to pdfplumber.
    """
    text = ""
    # Try PyMuPDF
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"PyMuPDF failed, falling back to pdfplumber: {e}")
        text = ""

    # Fallback to pdfplumber if text extraction is empty
    if not text.strip():
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber fallback failed: {e}")
            
    return text.strip()

def parse_docx(file_bytes: bytes) -> str:
    """
    Parses DOCX bytes using python-docx.
    """
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"python-docx parsing failed: {e}")
        
    return text.strip()

def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Main entry point for extracting text based on filename extension.
    """
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return parse_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")
